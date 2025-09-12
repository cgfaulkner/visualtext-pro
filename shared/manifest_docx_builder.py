#!/usr/bin/env python3
"""
Manifest-Based DOCX Review Builder
==================================

Builds DOCX review documents reading only from the ALT manifest.
NO LLaVA calls - all data comes from the manifest SSOT.

This eliminates the "missing ALT" problem by reading the actual ALT text
extracted from the PPTX rather than making assumptions based on thumbnails.
"""

from __future__ import annotations
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from alt_manifest import AltManifest

logger = logging.getLogger(__name__)
logger.info("LOG: injector_file=%s", __file__)


def generate_review_from_manifest(
    manifest_path: str,
    out_docx: str,
    title: Optional[str] = None,
    portrait: bool = True,
    run_id: str | None = None,
) -> str:
    """
    Generate ALT text review document from manifest (no LLaVA calls).
    
    Args:
        manifest_path: Path to alt_manifest.jsonl file
        out_docx: Output path for DOCX file
        title: Optional title for document
        portrait: If True, use portrait layout
        run_id: Optional identifier for tracking pipeline runs
        
    Returns:
        Path to generated DOCX file
    """
    logger.info("RUN_ID=%s manifest=%s", run_id, manifest_path)
    logger.info("Building DOCX review from manifest: %s", out_docx)
    
    # Load manifest
    try:
        manifest = AltManifest(Path(manifest_path))
        entries = manifest.get_all_entries()
        
        if not entries:
            logger.warning("No entries found in manifest")
            
    except Exception as e:
        logger.error(f"Failed to load manifest: {e}")
        raise RuntimeError(f"Could not load manifest file: {e}")
    
    # Create document
    doc = Document()
    
    # Setup layout
    if portrait:
        _setup_portrait_layout(doc)
    else:
        _setup_landscape_layout(doc)
    
    # Generate content
    doc_title = title or "ALT Text Review"
    _add_header_footer(doc, doc_title, Path(out_docx).name)
    _add_title_and_summary(doc, doc_title, entries, manifest)
    
    if entries:
        table = _create_review_table(doc, entries, portrait)
    else:
        doc.add_paragraph("No images found to review.")
    
    # Save document
    doc.save(out_docx)
    logger.info(f"DOCX review document saved: {out_docx}")
    
    # Log final statistics for verification
    _log_review_statistics(entries, manifest)
    
    return out_docx


def _setup_portrait_layout(doc):
    """Setup portrait page layout."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)


def _setup_landscape_layout(doc):
    """Setup landscape page layout."""
    section = doc.sections[0] 
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)


def _add_title_and_summary(doc, title: str, entries, manifest: AltManifest):
    """Add document title and summary statistics."""
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    _set_font_properties(title_run, size=16, bold=True)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(12)
    
    # Summary statistics from manifest
    stats = manifest.get_statistics()
    total_entries = stats['total_entries']
    with_current = stats['with_current_alt']
    with_suggested = stats['with_suggested_alt']
    
    summary_para = doc.add_paragraph()
    summary_text = f"Total images: {total_entries} | "
    summary_text += f"Current ALT text: {with_current} | "
    summary_text += f"Suggested ALT text: {with_suggested} | "
    summary_text += f"LLaVA calls made: {stats['llava_calls_made']}"
    
    summary_run = summary_para.add_run(summary_text)
    _set_font_properties(summary_run, size=10, italic=True, color="666666")
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_para.space_after = Pt(16)
    
    # Add processing mode summary
    source_breakdown = f"Existing: {stats['source_existing']}, Generated: {stats['source_generated']}, Cached: {stats['source_cached']}"
    breakdown_para = doc.add_paragraph()
    breakdown_run = breakdown_para.add_run(f"Sources: {source_breakdown}")
    _set_font_properties(breakdown_run, size=9, italic=True, color="888888")
    breakdown_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    breakdown_para.space_after = Pt(16)


def _create_review_table(doc, entries, portrait: bool):
    """Create the main review table from manifest entries."""
    # Sort entries by slide and image number
    sorted_entries = sorted(entries, 
                           key=lambda e: (e.slide_idx, e.image_number))
    
    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set fixed table layout for portrait
    if portrait:
        _setup_fixed_table_layout(table)
    
    # Setup header
    header_cells = table.rows[0].cells
    headers = ["Slide / Img", "Type", "Thumbnail", "Current ALT Text", "Suggested ALT Text", "Decorative"]
    
    for i, (cell, header) in enumerate(zip(header_cells, headers)):
        cell.text = header
        _format_header_cell(cell)
        if i == 2 and portrait:  # Thumbnail column padding (now column 2)
            _set_cell_margins(cell, right=180)
    
    _repeat_header_row(table.rows[0])
    
    # Add data rows
    for idx, entry in enumerate(sorted_entries, start=1):
        row = table.add_row()
        cells = row.cells
        
        # Slide/Image numbers  
        cells[0].text = f"{entry.slide_number}/{entry.image_number}"
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Shape Type (consistent labeling with PPT)
        shape_type_display = _get_shape_type_display(entry.shape_type, entry.is_group_child)
        cells[1].text = shape_type_display
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _format_shape_type_cell(cells[1], entry.shape_type)
        
        # Thumbnail (only for pictures)
        _add_thumbnail_to_cell(cells[2], entry, portrait)
        
        # Current ALT Text (from manifest - what was in PPTX originally)
        current_alt = (
            entry.existing_alt
            or entry.current_alt
            or (entry.final_alt if entry.decision_reason == "preserved" else "")
        ).strip()
        if idx <= 5:
            logger.debug(
                "existing_alt=%r docx_current=%r key=%s",
                entry.existing_alt,
                current_alt,
                entry.key,
            )
        cells[3].text = current_alt if current_alt else "[No ALT text]"
        _format_alt_text_cell(cells[3], bool(current_alt))
        
        # Suggested ALT Text (from manifest - final_alt is the SSOT)
        final_alt = (entry.final_alt or entry.suggested_alt).strip()
        
        # Apply synchronization logic based on decision_reason
        if entry.decision_reason == "preserved" or (entry.had_existing_alt and entry.source == "existing"):
            # Preserve mode: Current ALT = Suggested ALT
            suggested_label = f"{current_alt} (existing ALT preserved)" if current_alt else "[Preservation error]"
            cells[4].text = suggested_label
        elif entry.decision_reason in ["generated", "shape_fallback"] or final_alt:
            # Generated ALT text
            cells[4].text = final_alt if final_alt else "[Generation failed]"
        else:
            # No ALT text available
            cells[4].text = "[Generate needed]"
        
        _format_alt_text_cell(cells[4], bool(final_alt or current_alt))
        
        # Decorative checkbox
        is_decorative = _is_decorative_heuristic(entry)
        cells[5].text = "☑" if is_decorative else "☐"
        cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set row properties
        _set_row_no_break(row)
    
    return table


def _setup_fixed_table_layout(table):
    """Setup fixed table layout with portrait column widths."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    # Set column widths for portrait (adjusted for 6 columns)
    col_widths = [
        Inches(0.7),    # Slide/Img
        Inches(0.8),    # Type
        Inches(1.2),    # Thumbnail
        Inches(1.8),    # Current ALT
        Inches(1.8),    # Suggested ALT
        Inches(0.7)     # Decorative
    ]
    
    for i, width in enumerate(col_widths):
        if i < len(table.columns):
            table.columns[i].width = width


def _add_thumbnail_to_cell(cell, entry, portrait: bool):
    """Add thumbnail image to cell if available."""
    thumbnail_path = entry.thumbnail_path
    
    if thumbnail_path and Path(thumbnail_path).exists():
        try:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            # Size thumbnail appropriately for column width
            thumb_width = Inches(1.2) if portrait else Inches(1.5)
            run.add_picture(thumbnail_path, width=thumb_width)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            logger.debug(f"Could not add thumbnail for {entry.key}: {e}")
            cell.text = "[Image]"
    else:
        cell.text = "[No preview]"
    
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if portrait:
        _set_cell_margins(cell, right=180)  # Match header padding


def _get_shape_type_display(shape_type: str, is_group_child: bool) -> str:
    """
    Get consistent display label for shape type, matching PPT injector labels.
    """
    if is_group_child:
        prefix = "Group child: "
    else:
        prefix = ""
    
    # Map internal shape types to user-friendly display names
    type_mapping = {
        "PICTURE": "Image",
        "AUTO_SHAPE": "PowerPoint shape",
        "TEXT_BOX": "Text box",
        "LINE": "Line",
        "CONNECTOR": "Connector line", 
        "TABLE": "Table",
        "GROUP": "Group parent",
    }
    
    display_name = type_mapping.get(shape_type, shape_type.replace('_', ' ').title())
    return f"{prefix}{display_name}"


def _format_shape_type_cell(cell, shape_type: str):
    """Format shape type cell with appropriate styling."""
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    
    # Different styling based on shape type
    if shape_type == "PICTURE":
        _set_font_properties(run, size=9, bold=True, color="0066CC")
    else:
        _set_font_properties(run, size=9, color="666666")
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _is_decorative_heuristic(entry) -> bool:
    """Simple heuristic to suggest if image might be decorative."""
    # Very small images might be decorative
    if entry.width_px > 0 and entry.height_px > 0:
        if entry.width_px < 50 or entry.height_px < 50:
            return True
    
    # No ALT text provided might suggest decorative intent
    if not entry.current_alt and not entry.suggested_alt:
        return True
        
    return False


def _format_header_cell(cell):
    """Format header cell styling."""
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(cell.text)
    _set_font_properties(run, size=11, bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _shade_cell(cell, "E7E7E7")


def _format_alt_text_cell(cell, has_content: bool):
    """Format ALT text cell based on content."""
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    
    if has_content:
        _set_font_properties(run, size=10)
    else:
        _set_font_properties(run, size=10, italic=True, color="999999")
    
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _add_header_footer(doc, title: str, filename: str):
    """Add professional header and footer."""
    for section in doc.sections:
        # Header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.clear()
        
        title_run = header_para.add_run(f"{title} - Manifest Review")
        _set_font_properties(title_run, size=11, bold=True)
        
        header_para.add_run('\t' * 6)
        date_run = header_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
        _set_font_properties(date_run, size=10, color="666666")
        
        header_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
        
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.clear()
        
        filename_run = footer_para.add_run(filename)
        _set_font_properties(filename_run, size=10, color="666666")
        
        footer_para.add_run('\t' * 6)
        page_run = footer_para.add_run("Page ")
        _set_font_properties(page_run, size=10, color="666666")
        
        footer_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))


def _log_review_statistics(entries, manifest: AltManifest):
    """Log final statistics for verification."""
    stats = manifest.get_statistics()
    
    logger.info("📊 DOCX Review Statistics:")
    logger.info(f"   Total entries in manifest: {stats['total_entries']}")
    logger.info(f"   Current ALT text found: {stats['with_current_alt']}")
    logger.info(f"   Suggested ALT text available: {stats['with_suggested_alt']}")
    logger.info(f"   Source breakdown:")
    logger.info(f"     - Existing (preserved): {stats['source_existing']}")
    logger.info(f"     - Generated (LLaVA): {stats['source_generated']}")
    logger.info(f"     - Cached (reused): {stats['source_cached']}")
    logger.info(f"   LLaVA calls made: {stats['llava_calls_made']}")
    
    # This should now show correct counts instead of "Missing current ALT: X"


# Utility functions (same as before but included for completeness)

def _set_font_properties(run, size=11, bold=False, italic=False, color=None):
    """Set font properties for a run."""
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_margins(cell, left=None, right=None, top=None, bottom=None):
    """Set cell margins in twips."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        if val is None:
            continue
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _shade_cell(cell, hex_color):
    """Apply background shading to a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _repeat_header_row(row):
    """Make table header row repeat on every page."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def _set_row_no_break(row):
    """Prevent row from breaking across pages."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)
