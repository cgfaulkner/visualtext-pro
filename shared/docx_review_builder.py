#!/usr/bin/env python3
"""
DOCX Review Builder - Clean Pipeline Approach
=============================================

Builds DOCX review documents using the clean pipeline artifacts:
- visual_index.json: For thumbnails/labels/metadata
- current_alt_by_key.json: For Current ALT Text column  
- final_alt_map.json: For Suggested ALT Text column

No hidden reads from thumbnails or other sources - explicit JSON inputs only.
"""

from __future__ import annotations
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, Optional, Sequence

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from pipeline_artifacts import normalize_final_alt_map


logger = logging.getLogger(__name__)


STATUS_LABELS = {
    'preserve_existing': 'Preserve existing',
    'use_generated': 'Use generated',
    'no_alt_available': 'No ALT available',
    'flag_for_review': 'Needs review',
    'needs_review': 'Needs review',
    'custom_alt': 'Custom ALT',
    'manual_update': 'Manual update',
    'mark_decorative': 'Decorative',
    'decorative': 'Decorative',
}


def generate_alt_review_doc(
    visual_index_path: str,
    current_alt_by_key_path: str,
    final_alt_map_path: str,
    out_docx: str,
    portrait: bool = True,
    title: Optional[str] = None,
    config_manager = None
) -> str:
    """
    Generate ALT text review document from clean pipeline artifacts.
    
    Args:
        visual_index_path: Path to visual_index.json from Phase 1
        current_alt_by_key_path: Path to current_alt_by_key.json from Phase 1 (optional).
            If not provided or file is missing, current ALT text will be derived from
            visual_index entries (current_alt, existing_alt, or current_alt_text fields).
        final_alt_map_path: Path to final_alt_map.json from Phase 3
        out_docx: Output path for DOCX file
        portrait: If True, use portrait layout with fixed column widths
        title: Optional title for the document
        
    Returns:
        Path to generated DOCX file
    """
    logger.info(f"Building DOCX review document: {out_docx}")
    
    # Load required JSON artifacts
    try:
        with open(visual_index_path, 'r', encoding='utf-8') as f:
            visual_index = json.load(f)
            
        with open(final_alt_map_path, 'r', encoding='utf-8') as f:
            final_alt_map = json.load(f)

        final_alt_map = normalize_final_alt_map(final_alt_map)

    except Exception as e:
        logger.error(f"Failed to load required pipeline artifacts: {e}")
        raise RuntimeError(f"Could not load required JSON files: {e}")
    
    # Load optional current_alt_by_key.json (legacy file)
    # If missing, will be derived from visual_index below
    current_alt_by_key = {}
    try:
        with open(current_alt_by_key_path, 'r', encoding='utf-8') as f:
            current_alt_by_key = json.load(f)
        logger.debug(f"Loaded current_alt_by_key from {current_alt_by_key_path}")
    except (FileNotFoundError, IOError, json.JSONDecodeError) as e:
        logger.debug(f"current_alt_by_key.json not found or invalid: {e}. Will derive from visual_index.")
    
    # Derive current_alt_by_key from visual_index if file wasn't loaded
    if not current_alt_by_key:
        current_alt_by_key = {}
        for instance_key, entry in visual_index.items():
            current_alt = (
                entry.get("current_alt")
                or entry.get("existing_alt")
                or entry.get("current_alt_text")
                or ""
            ).strip()
            if current_alt:
                current_alt_by_key[instance_key] = current_alt
    
    if not visual_index:
        logger.warning("No images found in visual_index")
        
    # Create document
    doc = Document()
    
    # Setup document properties
    if portrait:
        _setup_portrait_layout(doc)
    else:
        _setup_landscape_layout(doc)
    
    # Add header and title
    doc_title = title or "ALT Text Review"

    # Get mode from config manager if available
    mode_info = ""
    if config_manager:
        try:
            mode = config_manager.get_alt_mode()
            mode_info = f" (Mode: {mode})"
        except:
            pass

    _add_header_footer(doc, doc_title + mode_info, Path(out_docx).name)
    
    # Add document title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{doc_title}")
    _set_font_properties(title_run, size=16, bold=True)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(12)
    
    # Add summary
    summary_para = doc.add_paragraph()
    summary_text = f"Total images: {len(visual_index)} | "
    summary_text += f"With existing ALT: {len(current_alt_by_key)} | "
    final_coverage = sum(
        1
        for record in final_alt_map.values()
        if (record.get('final_alt') or record.get('existing_alt') or record.get('generated_alt'))
    )
    summary_text += f"Final coverage: {final_coverage}"
    summary_run = summary_para.add_run(summary_text)
    _set_font_properties(summary_run, size=10, italic=True, color="666666")
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_para.space_after = Pt(16)
    
    # Create table
    if portrait:
        table = _create_portrait_table(doc, visual_index, current_alt_by_key, final_alt_map)
    else:
        table = _create_landscape_table(doc, visual_index, current_alt_by_key, final_alt_map)
    
    # Save document
    doc.save(out_docx)
    logger.info(f"DOCX review document saved: {out_docx}")
    
    return out_docx


def _setup_portrait_layout(doc):
    """Setup portrait page layout with optimized margins."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)  # More space for header
    section.bottom_margin = Inches(0.75)


def _setup_landscape_layout(doc):
    """Setup landscape page layout."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)


def _create_portrait_table(
    doc,
    visual_index: Dict,
    current_alt_by_key: Dict,
    final_alt_map: Dict[str, Dict[str, Any]],
):
    """Create portrait-oriented table with focused review columns."""
    column_widths = (0.70, 1.50, 1.75, 1.75, 0.60, 0.50, 0.70)

    sorted_keys = sorted(
        visual_index.keys(),
        key=lambda k: (
            visual_index[k].get('slide_idx', 0),
            visual_index[k].get('image_number', 0),
        ),
    )

    table = doc.add_table(rows=1, cols=len(column_widths))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    width_measures = [Inches(width) for width in column_widths]
    for i, width in enumerate(width_measures):
        table.columns[i].width = width

    headers = [
        'Slide/ID',
        'Thumbnail',
        'Current',
        'Generated ALT',
        'Status',
        'Decor',
        'Notes',
    ]

    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        cell.text = header
        _format_header_cell(cell)

    _repeat_header_row(table.rows[0])

    for key in sorted_keys:
        visual_info = visual_index.get(key, {})
        row = table.add_row()
        cells = row.cells

        slide_num = visual_info.get('slide_number', '?')
        img_num = visual_info.get('image_number', '?')
        cells[0].text = f"{slide_num}/{img_num}"
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        thumbnail_path = visual_info.get('thumbnail_path')
        cells[1].text = ''
        paragraph = cells[1].paragraphs[0]
        run = paragraph.add_run()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            if thumbnail_path and Path(thumbnail_path).exists():
                run.add_picture(thumbnail_path, width=Inches(1.5))
            else:
                raise FileNotFoundError(thumbnail_path or 'Missing thumbnail')
        except Exception as exc:
            logger.debug(f"Could not add thumbnail for {key}: {exc}")
            cells[1].text = "[No preview]"
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        record = final_alt_map.get(key, {}) if isinstance(final_alt_map, dict) else {}
        if not isinstance(record, dict):
            record = {}

        current_alt = (current_alt_by_key.get(key) or '').strip()
        current_display = current_alt if current_alt else "[No ALT text]"
        cells[2].text = current_display
        _format_alt_text_cell(cells[2], bool(current_alt))
        _ensure_wrap(cells[2])
        _set_cell_margins(cells[2], left=80, right=80)

        generated_alt = _get_generated_alt_text(record, visual_info, current_alt)
        generated_display = generated_alt if generated_alt else "[Not generated]"
        cells[3].text = generated_display
        _format_alt_text_cell(cells[3], bool(generated_alt))
        _ensure_wrap(cells[3])
        _set_cell_margins(cells[3], left=80, right=80)

        status_text = _resolve_status_label(record, current_alt, generated_alt)
        cells[4].text = status_text
        cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        is_decorative = _is_decorative_image(
            visual_info,
            current_alt,
            generated_alt,
        )
        decision_value = (record.get('decision') or '').strip() if record else ''
        if decision_value in {'mark_decorative', 'decorative'}:
            is_decorative = True
        cells[5].text = '☑' if is_decorative else '☐'
        cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        notes_text = _build_notes_text(record, visual_info)
        cells[6].text = notes_text
        _format_alt_text_cell(cells[6], bool(notes_text))
        _ensure_wrap(cells[6])
        _set_cell_margins(cells[6], left=80, right=80)

        cells[6].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        _set_row_no_break(row)

    return table


def _create_landscape_table(
    doc,
    visual_index: Dict,
    current_alt_by_key: Dict,
    final_alt_map: Dict,
):
    """Create landscape-oriented table with wider ALT text columns."""
    landscape_widths = (0.85, 1.6, 1.2, 1.6, 1.2, 1.2, 0.6, 1.15)
    return _create_review_table(
        doc,
        visual_index,
        current_alt_by_key,
        final_alt_map,
        landscape_widths,
    )


def _create_review_table(
    doc,
    visual_index: Dict,
    current_alt_by_key: Dict,
    final_alt_map: Dict[str, Dict[str, Any]],
    column_widths: Sequence[float],
):
    """Build the shared review table for portrait or landscape layouts."""
    # Sort images by slide and image number for consistent ordering
    sorted_keys = sorted(
        visual_index.keys(),
        key=lambda k: (
            visual_index[k].get('slide_idx', 0),
            visual_index[k].get('image_number', 0),
        ),
    )

    table = doc.add_table(rows=1, cols=len(column_widths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set fixed table layout
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    width_measures = [Inches(width) for width in column_widths]
    for i, width in enumerate(width_measures):
        table.columns[i].width = width

    headers = [
        "Slide / Img",
        "Thumbnail",
        "Current ALT Text",
        "Suggested ALT Text",
        "Existing ALT",
        "Generated ALT",
        "Decorative",
        "Status",
    ]

    header_cells = table.rows[0].cells
    for i, (cell, header) in enumerate(zip(header_cells, headers)):
        cell.text = header
        _format_header_cell(cell)
        if i == 1:
            _set_cell_margins(cell, right=180)

    _repeat_header_row(table.rows[0])

    for key in sorted_keys:
        visual_info = visual_index[key]
        row = table.add_row()
        cells = row.cells

        slide_num = visual_info.get('slide_number', '?')
        img_num = visual_info.get('image_number', '?')
        cells[0].text = f"{slide_num}/{img_num}"
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        thumbnail_path = visual_info.get('thumbnail_path')
        if thumbnail_path and Path(thumbnail_path).exists():
            try:
                paragraph = cells[1].paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(thumbnail_path, width=Inches(1.2))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.debug(f"Could not add thumbnail for {key}: {e}")
                cells[1].text = "[Image]"
        else:
            cells[1].text = "[No preview]"
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_margins(cells[1], right=180)

        current_alt = current_alt_by_key.get(key, "").strip()
        cells[2].text = current_alt if current_alt else "[No ALT text]"
        _format_alt_text_cell(cells[2], bool(current_alt))

        record = (
            final_alt_map.get(key, {})
            if isinstance(final_alt_map, dict)
            else {}
        )
        if not isinstance(record, dict):
            record = {}
        suggested_alt = _select_suggested_alt(record)
        final_existing_alt = (record.get('existing_alt') or '').strip()
        final_generated_alt = (record.get('generated_alt') or '').strip()

        # Suggested ALT Text column: show generated/proposed ALT only, never current_alt
        display_suggested = suggested_alt if suggested_alt else "[Generate needed]"
        cells[3].text = display_suggested
        _format_alt_text_cell(cells[3], bool(suggested_alt))

        existing_display = (
            final_existing_alt if final_existing_alt else "[No ALT text]"
        )
        cells[4].text = existing_display
        _format_alt_text_cell(cells[4], bool(final_existing_alt))

        # Generated ALT column: show only final_generated_alt, never existing ALT
        generated_display = (
            final_generated_alt if final_generated_alt else "[Not generated]"
        )
        cells[5].text = generated_display
        _format_alt_text_cell(cells[5], bool(final_generated_alt))

        is_decorative = _is_decorative_image(visual_info, current_alt, suggested_alt)
        cells[6].text = "☐"
        cells[6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Status column: use same logic as portrait table
        generated_alt_for_status = _get_generated_alt_text(record, visual_info, current_alt)
        status_text = _resolve_status_label(record, current_alt, generated_alt_for_status)
        cells[7].text = status_text
        cells[7].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        _set_row_no_break(row)

    return table


def _clean_text(value: Any) -> str:
    """Return stripped text for the provided value."""
    if value is None:
        return ''
    return str(value).strip()


def _get_generated_alt_text(
    record: Dict[str, Any],
    visual_info: Dict[str, Any],
    current_alt: str,
) -> str:
    """Prefer enriched generated ALT text with graceful fallbacks."""
    generated_alt = _clean_text(record.get('generated_alt')) if record else ''
    if generated_alt:
        return generated_alt

    visual_generated = _clean_text(visual_info.get('generated_alt'))
    if visual_generated:
        return visual_generated

    final_alt = _clean_text(record.get('final_alt')) if record else ''
    if final_alt and final_alt != current_alt:
        return final_alt

    return ''


def _resolve_status_label(
    record: Dict[str, Any],
    current_alt: str,
    generated_alt: str,
) -> str:
    """Resolve a human-friendly status label with sensible defaults."""
    decision_value = _clean_text(record.get('decision')) if record else ''
    if decision_value:
        normalized = decision_value.lower()
        if normalized in STATUS_LABELS:
            return STATUS_LABELS[normalized]
        return decision_value.replace('_', ' ').strip().title()

    final_alt = _clean_text(record.get('final_alt')) if record else ''
    if final_alt:
        if generated_alt and final_alt == generated_alt:
            return STATUS_LABELS['use_generated']
        if current_alt and final_alt == current_alt:
            return STATUS_LABELS['preserve_existing']
        return STATUS_LABELS['custom_alt']

    if generated_alt and current_alt:
        return 'Existing + Generated'
    if generated_alt:
        return 'Generated available'
    if current_alt:
        return 'Existing only'

    return STATUS_LABELS['no_alt_available']


def _build_notes_text(
    record: Dict[str, Any],
    visual_info: Optional[Dict[str, Any]] = None,
) -> str:
    """Assemble provenance and notes information for display."""
    notes_parts: list[str] = []

    if record:
        source_existing = _clean_text(record.get('source_existing'))
        if source_existing:
            notes_parts.append(f"Existing: {source_existing}")

        source_generated = _clean_text(record.get('source_generated'))
        if source_generated:
            notes_parts.append(f"Generated: {source_generated}")

        record_notes = _clean_text(record.get('notes'))
        if record_notes:
            notes_parts.append(record_notes)

        provenance = record.get('provenance')
        if isinstance(provenance, dict):
            for key, value in provenance.items():
                provenance_value = _clean_text(value)
                if provenance_value:
                    key_text = _clean_text(key) or 'Provenance'
                    notes_parts.append(f"{key_text}: {provenance_value}")
        elif isinstance(provenance, list):
            for value in provenance:
                provenance_value = _clean_text(value)
                if provenance_value:
                    notes_parts.append(provenance_value)
        else:
            provenance_text = _clean_text(provenance)
            if provenance_text:
                notes_parts.append(provenance_text)

    if visual_info:
        visual_provenance = _clean_text(visual_info.get('provenance'))
        if visual_provenance:
            notes_parts.append(visual_provenance)

    return ' • '.join(notes_parts)


def _select_suggested_alt(record: Dict[str, Any]) -> str:
    """Return the preferred suggested ALT text from a final_alt_map record."""
    if not isinstance(record, dict):
        return ""

    final_alt = (record.get('final_alt') or "").strip()
    generated_alt = (record.get('generated_alt') or "").strip()

    return final_alt or generated_alt


def _format_header_cell(cell):
    """Format header cell with proper styling."""
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(cell.text)
    _set_font_properties(run, size=11, bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Add background shading
    _shade_cell(cell, "E7E7E7")  # Light gray


def _format_alt_text_cell(cell, has_content: bool):
    """Format ALT text cell based on whether it has content."""
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run()
        
    if has_content:
        _set_font_properties(run, size=10)
    else:
        _set_font_properties(run, size=10, italic=True, color="999999")
    
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _is_decorative_image(visual_info: Dict, current_alt: str, suggested_alt: str) -> bool:
    """
    Heuristic to determine if image might be decorative.
    This is just a hint - final determination should be manual.
    """
    # Check size (very small images might be decorative)
    width = visual_info.get('width_px', 0)
    height = visual_info.get('height_px', 0)
    if width > 0 and height > 0 and (width < 50 or height < 50):
        return True
    
    # Check if both ALT texts are empty (might indicate decorative)
    if not current_alt and not suggested_alt:
        return True
        
    return False


def _add_header_footer(doc, title: str, filename: str):
    """Add professional header and footer."""
    for section in doc.sections:
        # Header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.clear()
        
        # Left-aligned title
        title_run = header_para.add_run(f"{title}")
        _set_font_properties(title_run, size=11, bold=True)
        
        # Right-aligned timestamp
        header_para.add_run('\t' * 6)
        date_run = header_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
        _set_font_properties(date_run, size=10, color="666666")
        
        header_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
        
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.clear()
        
        # Left-aligned filename
        filename_run = footer_para.add_run(filename)
        _set_font_properties(filename_run, size=10, color="666666")
        
        # Right-aligned page numbers
        footer_para.add_run('\t' * 6)
        page_run = footer_para.add_run("Page ")
        _set_font_properties(page_run, size=10, color="666666")
        
        footer_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))


def _set_font_properties(run, size=11, bold=False, italic=False, color=None):
    """Set font properties for a run."""
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_margins(cell, left=None, right=None, top=None, bottom=None):
    """Set cell margins in twips (1/20 pt). 180 ≈ 0.125"."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        if val is None:
            continue
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _shade_cell(cell, hex_color):
    """Apply background shading to a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _repeat_header_row(row):
    """Make table header row repeat on every page."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def _set_row_no_break(row):
    """Prevent row from breaking across pages."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def _ensure_wrap(cell) -> None:
    """Ensure cell content can wrap across lines."""
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap_tag = qn('w:noWrap')
    for no_wrap in list(tc_pr.findall(no_wrap_tag)):
        tc_pr.remove(no_wrap)

    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.keep_with_next = False
